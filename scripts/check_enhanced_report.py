from docx import Document

doc = Document('reports/FINAL_COMPREHENSIVE_REPORT_ENHANCED.docx')

print('='*60)
print('ENHANCED REPORT STATISTICS')
print('='*60)
print(f'\nTotal Paragraphs: {len(doc.paragraphs)}')
print(f'Total Tables: {len(doc.tables)}')

figures = [p.text for p in doc.paragraphs if 'Figure' in p.text and ':' in p.text]
print(f'Total Figures Referenced: {len(figures)}')

print(f'\n{"="*60}')
print('FIGURE LIST (First 15)')
print('='*60)
for i, fig in enumerate(figures[:15], 1):
    print(f'{i}. {fig[:100]}')

print(f'\n{"="*60}')
print('DASHBOARD SECTIONS')
print('='*60)
dashboard_refs = [p.text for p in doc.paragraphs if 'Dashboard' in p.text and ('Screenshot' in p.text or 'Visual' in p.text)]
for i, ref in enumerate(dashboard_refs[:10], 1):
    print(f'{i}. {ref[:100]}')

print(f'\n{"="*60}')
print('CORRELATION/ASSOCIATION ANALYSIS')
print('='*60)
correlation_refs = [p.text for p in doc.paragraphs if 'correlation' in p.text.lower() or 'association' in p.text.lower()]
for i, ref in enumerate(correlation_refs[:5], 1):
    print(f'{i}. {ref[:100]}')

print(f'\n{"="*60}')
print('REPORT ENHANCEMENT SUMMARY')
print('='*60)
print(f'✓ Total Paragraphs: {len(doc.paragraphs)} (was ~387)')
print(f'✓ Visual References: {len(figures)} figures')
print(f'✓ Dashboard Screenshots: {len([f for f in figures if "Dashboard" in f])} sections')
print(f'✓ Association Matrix: {"Yes" if any("correlation" in f.lower() or "matrix" in f.lower() for f in figures) else "No"}')
print('='*60)
